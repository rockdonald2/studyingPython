""" def div42by(divideBy):
    try:
        return 42 / divideBy
    except ZeroDivisionError:
        print("Error, skip current call.")


print(div42by(0))

print('How many cats do you have?')
numCats = input()
try:
    if int(numCats) >= 2:
        print("That's a lot of cats")
    else:
        print("That's not that many cats")
except ValueError:
    print('Please enter a digit, not a transcript')


# Guess a number game
import random
print('I am thinking of a number between 1 and 50, you think you can guess it easily right? Then, let\'s play!')
randomNumber = random.randint(1, 50)

for guessesTaken in range(1, 11):
    print('Take a guess')
    while True:
        try:
            guess = int(input())
            break
        except ValueError:
            print('Please enter a digit')
            continue

    if guess < randomNumber:
        print('Your guess is too low')
    elif guess > randomNumber:
        print('Your guess is too damn high')
    else:
        break  # current guess

if guess == randomNumber:
    print('Damn, you nailed it, good job!')
else:
    print('Damn')



 """

""" Regular expressions """

""" message = 'Call me 415-555-1011 tomorrow'
phoneNumRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d')
mo = phoneNumRegex.search(message)
print(mo.group()) """

""" phoneNumRegex = re.compile(r'(\d\d\d)-(\d\d\d)-(\d\d\d)')
mo = phoneNumRegex.search('My number is 415-555-1111')
print(mo.group(1, 2)) """

""" phoneNumRegex = re.compile(r'\(\d\d\d\) \d\d\d-\d\d\d')
mo = phoneNumRegex.search('My number is (415) 555-1111')
print(mo.group()) """

""" batRegex = re.compile(r'Bat(man|mobile|copter|bat)')
mo = batRegex.search('Batmobile lost a wheel')
print(mo.group()) """

""" batRegex = re.compile(r'Bat(wo)?man') # the group can appear zero or one times in order to match the pattern
mo = batRegex.search('The adventures of Batwoman')
print(mo.group()) """

""" phoneNumRegex = re.compile(r'(\d\d\d-)?\d\d\d-\d\d\d\d')
mo = phoneNumRegex.search('My phone number is 415-555-1111')
print(mo.group())

batRegex = re.compile(r'Bat(wo)*man') # zero or more
mo2 = batRegex.search('Batwowowowoman')
print(mo2.group())

batRege2 = re.compile(r'Bat(wo)+man') # one or more
mo3 = batRege2.search('Batwowowoman')
print(mo3.group()) """

""" regEx = re.compile(r'(Ha){3}') # exactly 3 times
mo = regEx.search('HaHaHa')
print(mo.group()) """

""" regEx = re.compile(r'(ha){2,4}') # at least 2, at most 4 times
mo = regEx.search('hahaha')
print(mo.group()) """

""" regEx = re.compile(r'(\d){3,5}?')  # ? marks means a non-greedy match, which means if there are more matches than 5 the first 3 will be matched, greedy the longest match, non-greedy means the shortest possible match
mo = regEx.search('1234567890')
print(mo.group()) """

"""
phoneRegEx = re.compile(r'(\(\d\d\d\)) (\d\d\d)-(\d\d\d\d)')
print(phoneRegEx.findall('''Ivy Haddington
Los Angeles California • (123) 456-7891
ihaddington@email.com

SUMMARY
Energetic, creative Nanny with more than three years of experience in transporting children to and from activities, keeping play spaces neat and tidy and assisting with homework and projects.

EDUCATION
Coral Springs University
Aug '10 - May '14
Early Childhood Education/Wellness & Health Education

EXPERIENCE
Cloud Clearwater, Nanny
Feb '15 - Current
Work as a unit with families to make sure children receive care specially tailored to their needs
Ensure the safety and well-being of children at all times
Clean and sanitize all toys and equipment before and after use
Prepare delicious, nutritional and kid-friendly meals with attention to dietary restrictions and preferences
River Tech, Child Care Provider
Current - Current
Picked up and dropped off children from day cares, nurseries and school
Provided assistance with homework and creative projects, improving overall grades at school by 45% in one calendar year
Supervised play dates and fun trips to places like the park, swimming pool and zoo
SKILLS
CPR
Typing''')) """

lyrics = '''On the first day of Christmas my true love gave to me,
A partridge in a pear tree.
On the second day of Christmas my true love gave to me,
Two turtle doves and a partridge in a pear tree.
On the fourth day of Christ
On the third day of Christmas my true love gave to me, five golden rings
Four calling birds, three french hens,
Two turtle doves and a partridge in a pear tree
On the ninth day of Christmas, (on the eighth day of Christmas)
My true love gave to me (my true love gave to me)
Nine ladies dancing (eight maids a milking)
Eight ladies dancing (seven maids a milking)
Seven ladies dancing (six maids a milking)
Six ladies dancing (five)
Better not shout, you better not cry,
You better not in a pear tree
On the ninth, no!
On the eighth, come on!
On the seventh day of Christmas (deck the halls)
My true love gave to me (with boughs of holly)
Seven swans a swimming (here we come a wassailing)
Six geese a laying (among the leaves so)
Five golden rings (Fa-la-la-la-la)
Four calling birds
Three french hens
Two turtle doves (here we come a wassailing)
The boar's head in hand bear I (among the leaves so green)
Bedeck'd with bays
And partridge in a pear tree
On the eleventh day of Christmas my true love gave to me
Eleven pipers piping (ding)
Ten lords a leaping (dong)
Nine ladies dancing (ding)
Eight maids a milking (dong)
Seven swans a swimming (ding)
Six geese a laying (dong)
Five golden rings
Four calling birds
Three french hens
Two turtle doves
And Rudolph the red nosed reindeer!
On the twelfth day of Christmas my true love gave to me
I have a little dreidel, I made it out of clay,
And when it's dry and ready, my dreidel I shall play
Oh dreidel dreidel dreidel (hey, 12 days of Christmas)
(what? 8 days of Hanukkah. It's a Christmas medley)
On the twelfth day of Christmas my true love gave to me
Do do do do doot do (do de do de do de do de do de)
Do do do do doot do (do de do de do de do de do de)
On the twelfth day my true love gave to me
Twelve drummers drumming like Olympus above the Serengeti
Ba ba ba ba ba ba ba oh
Eleven pipers piping, ten lords a leaping (Ba ba ba ba ba ba ba oh)
Nine ladies dancing they were dancing for me
Eight maids a milking they were milking just for me
I had Christmas down in Africa
I had Christmas down in Africa
I had Christmas down in Africa (five golden rings)
I had Christmas down in Africa (five golden rings)
I had Christmas down in Africa (five golden rings)
Couldn't take the halls into the things we never had
Do do do do doot do (do de do de do de do de do de)
Do do do do doot do (do de do de do de do de do de)
Partridge in a big pear tree (do de do de do de do de do de)
Partridge in a big pear tree'''

"""xmasRegex = re.compile(r'\d+\s\w+')
print(xmasRegex.findall(lyrics))

vowelRegex = re.compile(r'[^aeiouAEIOU]')  # same as r'(a|e|i|o|u)'
print(vowelRegex.findall(lyrics))"""

"""beginsWithHello = re.compile(r'^Hello')
print(beginsWithHello.search('Hello there young man'))

endsWithHello = re.compile(r'hello$')
print(endsWithHello.search('There hello'))

allDigitsRegEx = re.compile(r'^\d+$')
print(allDigitsRegEx.search('22222'))

dotStar = re.compile(r'.*', re.DOTALL)
print(dotStar.search('Serve the public\nSlay the dragons'))

vowelRegex = re.compile(r'[aeiou]', re.IGNORECASE)
print(vowelRegex.findall('ManiAc killer'))"""

"""namesRegex = re.compile(r'Agent \w+')
print(namesRegex.findall('Agent Alice gave the secret documents to Agent Bob'))
print(namesRegex.sub('REDACTED', 'Agent Alice gave the secret documents to Agent Bob'))"""

"""namesRegex = re.compile(r'Agent (\w)\w*')
print(namesRegex.findall('Agent Alice gave the secret documents to Agent Bob'))
print(namesRegex.sub(r'Agent \1****', 'Agent Alice gave the secret documents to Agent Bob'))"""

"""regExObj = re.compile(r'''
\d\d\d # comment here
-
\d\d\d # comment there
-
\d\d\d\d
''', re.VERBOSE)
print(regExObj.search('415-555-1011'))"""

"""

import pyperclip, re


# Create a regex for phone numbers

phoneRegex = re.compile(r'''
(((\d\d\d) | (\(\d\d\d\)))?    # optional area code
(\s | -)    # first separator
(\d\d\d)    # first three digits
(-)    # second separator
(\d\d\d\d)    # four digits
((ext(\.)?\s |x)  (\d{2,5}))?)    # optional extension
''', re.VERBOSE)  # 415-555-1111, 555-1111, (415) 555-1111, 555-1111 ext 12345, ext.
# x12345

# Regex object for e-mail addresses

emailRegex = re.compile(r'''
# some.+_thing@something.com

[a-zA-Z0-9_.+]+    # name part
@   # @ symbol
[a-zA-Z0-9_.+]+    # domain name

''', re.VERBOSE | re.IGNORECASE)

# Get the text off the clipboard
text = pyperclip.paste()

# Extract the e-mail addresses and phone numbers of the phone directory
extractedPhone = phoneRegex.findall(text)
extractedEmail = emailRegex.findall(text)

allPhoneNumbers = []
for i in extractedPhone:
    allPhoneNumbers.append(i[0])

# Copy the extracted email/phone to the clipboard
results = '\n'.join(allPhoneNumbers) + '\n'.join(extractedEmail)
pyperclip.copy(results)

"""

"""import os

totalsize = 0
for filename in os.listdir('d:\\GitHub\\studyingSQL'):
    if not os.path.isfile(os.path.join('d:\\GitHub\\studyingSQL', filename)):
        continue
    totalsize = totalsize + os.path.getsize(os.path.join('d:\\GitHub\\studyingSQL', filename))

print(totalsize)"""

"""import shelve
shelfFile = shelve.open('mydata')
shelfFile['cats'] = ['Zophie', 'Pooka', 'Simon', 'Fat-tail', 'Cleo']
shelfFile.close()"""

"""import shutil
shutil.copy('d:\\GitHub\\studyingSQL\\modul11.sql', 'd:\\modul11copy.sql')
shutil.copytree('d:\\GitHub\\studyingSQL', 'd:\\modulbackup')"""

"""import send2trash
send2trash.send2trash()"""

"""import os
for folderName, subFolders, fileName in os.walk('d:\\GitHub'):
    print('The folder is ' + folderName)
    print('The subfolders in' + folderName + ' are: ' + str(subFolders))
    print('And the files are: ' + str(fileName))"""

"""import traceback

def boxPrint(symbol, width, height):
    if len(symbol) != 1:
        errorFile = open('error_log.txt', 'a')
        errorFile.write(traceback.format_exc())
        errorFile.close()
        raise Exception('Symbol argument has to be length of 1.')

    if width < 2 or height < 2:
        errorFile = open('error_log.txt', 'a')
        errorFile.write(traceback.format_exc())
        errorFile.close()
        raise Exception('Width or height must be greater than 2.')

    assert len(symbol) == 1 and width > 2 and height > 2, 'Bug raised'

    print(symbol * width)

    for i in range(height - 2):
        print(symbol + (' ' * (width - 2)) + symbol)

    print(symbol * width)


boxPrint('**', 15, 5)"""

"""import logging

logging.basicConfig(filename='error_log.txt', level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.disable(logging.WARNING)

logging.debug('Start of program')


def factorial(n):
    logging.debug('Start of factorial(%s)' % (n))
    total = 1
    for i in range(1, n + 1):
        total *= i
        logging.debug('i is %s, total is %s' % (i, total))

    logging.debug('Return value is %s' % total)
    return total


print(factorial(5))"""

"""import webbrowser

webbrowser.open('https:\\google.com')"""

"""import webbrowser, pyperclip, sys

# Check if command line arugments were passed

if len(sys.argv) > 1:
    # ['mapit.py', '870', 'Valencia', 'St.'] // space delimited -> '870 Valencia St.'
    address = ' '.join(sys.argv[1:])
else:
    # arguments are in clipboard
    address = pyperclip.paste()

# https://www.google.com/maps/place/<ADDRESS>
webbrowser.open('https://www.google.com/maps/place/' + address)"""

"""import requests
res = requests.get('https://automatetheboringstuff.com/files/rj.txt')
print(res.status_code)

print(res.text[:500])

res.raise_for_status()
badRes = requests.get('https://automatetheboringstuff.com/files/rj.txtaaaaaa')
badRes.raise_for_status()
playFile = open('RomeoAndJuliet.txt', 'wb')

for chunk in res.iter_content(100000):
    playFile.write(chunk)

playFile.close()"""

"""import bs4, requests

res = requests.get('https://www.emag.ro/365-mesaje-pentru-o-supermama-5948489351549/pd/DWFTRBBBM/?X-Search-Id=6aed704ba4ac530c7941&X-Product-Id=1129345&X-Search-Page=1&X-Search-Position=0&X-Section=search&X-MB=0&X-Search-Action=view')
res.raise_for_status()

soup = bs4.BeautifulSoup(res.text, 'html.parser')
elems = soup.select('#page-skin > div.container > div > div:nth-child(2) > div.col-sm-5.col-md-7.col-lg-7 > div > div > div.col-sm-12.col-md-6.col-lg-5 > form > div.product-highlight.product-page-pricing > p.product-new-price')
print(elems[0].text.strip())"""


"""from selenium import webdriver
browser = webdriver.Chrome('C:\Program Files (x86)\Google\Chrome\chromedriver.exe')
browser.get('https:\\google.com')
elem = browser.find_element_by_css_selector('#tsf > div:nth-child(2) > div.A8SBwf > div.RNNXgb > div > div.a4bIc > input')
print(elem)
elem.click()
elems = browser.find_elements_by_css_selector('a')
print(elems)
elem.send_keys('smart')
elem.submit()"""

"""import smtplib
conn = smtplib.SMTP('smtp.gmail.com', 587)
print(conn)
print(conn.ehlo())
conn.starttls() # begin encryption
print(conn.login('rockdonald2@gmail.com', ''))
conn.sendmail('rockdonald2@gmail.com', 'rockdonald2@gmail.com', 'Subject: So long...\n\nDear Zsolt, \nSo long, and thanks for all the fish.\n\n-Zsolt')
conn.quit()"""

"""import imapclient, pyzmail

conn = imapclient.IMAPClient('imap.gmail.com', ssl=True)
print(conn.login('rockdonald2@gmail.com', ''))
select_info = conn.select_folder('INBOX', readonly=True)
UIDs = conn.search(['FROM', 'rockdonald2@gmail.com'])
print(UIDs)
rawMessage = conn.fetch([9117], ['BODY[]', 'FLAGS'])
message = pyzmail.PyzMessage.factory(rawMessage[9117][b'BODY[]'])

print(message.get_subject())
print(message.get_addresses('from'))
print(message.get_addresses('bcc'))
print(message.text_part)
print(message.html_part)
print(message.text_part.get_payload().decode('UTF-8'))"""

"""import openpyxl, os

os.chdir('c:\\users\\Lukács Zsolt\\Downloads')
print(os.getcwd())

workbook = openpyxl.load_workbook('example.xlsx')
sheet = workbook['Sheet1']
cell = sheet['B2']
print(cell.value)

print(sheet.cell(row=1, column=2))

for i in range(1, 8):
    print(i, sheet.cell(row=i, column=2).value)


wb = openpyxl.Workbook()
print(wb.sheetnames)
sh = wb['Sheet']
sh['A1'] = 42
sh['A2'] = 'Hello'
sh2 = wb.create_sheet('Sheet2')
sh2.title = 'Idunno'
wb.save('example2.xlsx')"""


"""import PyPDF2, os

os.chdir('c:\\users\\Lukács Zsolt\\Downloads')

pdfFile = open('meetingminutes1.pdf', 'rb')

reader = PyPDF2.PdfFileReader(pdfFile)
print(reader)
print(reader.numPages)
print(reader.getPage(15).extractText())"""

"""import PyPDF2, os

os.chdir('c:\\Users\\Lukács Zsolt\\Downloads')
pdf1File = open('meetingminutes1.pdf', 'rb')
pdf2File = open('meetingminutes2.pdf', 'rb')
reader1 = PyPDF2.PdfFileReader(pdf1File)
reader2 = PyPDF2.PdfFileReader(pdf2File)

writer = PyPDF2.PdfFileWriter()

for pageNum in range(reader1.numPages):
    page = reader1.getPage(pageNum)
    writer.addPage(page)

for pageNum in range(reader2.numPages):
    page = reader2.getPage(pageNum)
    writer.addPage(page)

outputFile = open('combinedminutes.pdf', 'wb')
writer.write(outputFile)
outputFile.close()
pdf1File.close()
pdf2File.close()"""

"""import docx

d = docx.Document('c:\\users\\Lukács Zsolt\\Downloads\\demo.docx')
print(d)

print(d.paragraphs[1].text)

p = d.paragraphs[1]
print(p.runs[1].text)
print(p.runs[1].bold)
print(p.runs[3].italic)
p.runs[2].underline = True

d.save('c:\\users\\Lukács Zsolt\\Downloads\\demo2.docx')

d2 = docx.Document()
d2.add_paragraph('Hello this is a paragraph written to a doc file.')
d2.add_paragraph('Another paragraph.')
d2.save('c:\\users\\Lukács Zsolt\\Downloads\\demo3.docx')

p2 = d2.paragraphs[0]
p2.add_run('This is a new run.')
print(p2.runs)
p2.runs[1].bold = True
d2.save('c:\\users\\Lukács Zsolt\\Downloads\\demo3.docx')"""


"""import docx


def getText(filename):  # returns a single string value
    doc = docx.Document(filename)
    fullText = []
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
    return '\n'.join(fullText)


print(getText('c:\\users\\Lukács Zsolt\\Downloads\\demo.docx'))"""


import pyautogui